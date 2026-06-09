from __future__ import annotations

from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "outputs" / "script-analysis-report.md"
OUTPUT = ROOT / "outputs" / "script-analysis-report.docx"


def clean_cell(value: str) -> str:
    value = value.replace(r"\|", "|")
    value = strip_inline_markdown(value)
    value = re.sub(r"\s+", " ", value).strip()
    return value


def strip_inline_markdown(value: str) -> str:
    value = re.sub(r"\*\*([^*]+)\*\*", r"\1", value)
    value = re.sub(r"`([^`]+)`", r"\1", value)
    return value


def split_md_row(line: str) -> list[str]:
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    cells: list[str] = []
    cur = []
    escaped = False
    for char in line:
        if escaped:
            cur.append(char)
            escaped = False
            continue
        if char == "\\":
            escaped = True
            cur.append(char)
            continue
        if char == "|":
            cells.append(clean_cell("".join(cur)))
            cur = []
        else:
            cur.append(char)
    cells.append(clean_cell("".join(cur)))
    return cells


def is_separator(line: str) -> bool:
    return bool(re.match(r"^\s*\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?\s*$", line))


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = "D7DEE8") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:color"), color)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_width(cell, width_cm: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_cm * 567)))
    tc_w.set(qn("w:type"), "dxa")


def style_run(run, size: float = 8, bold: bool = False, color: str | None = None) -> None:
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_paragraph(document: Document, text: str, style: str | None = None) -> None:
    para = document.add_paragraph(style=style)
    run = para.add_run(text)
    style_run(run, 9.5 if style != "Title" else 24, bold=style in {"Title", "Heading 1", "Heading 2"})
    if style == "Title":
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run.font.color.rgb = RGBColor(31, 78, 121)


def add_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    columns = max(len(row) for row in rows)
    table = document.add_table(rows=1, cols=columns)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"

    if columns == 3:
        widths = [5.0, 12.0, 13.0]
        font_size = 8.5
    elif columns == 9:
        widths = [5.0, 2.4, 5.7, 2.6, 3.4, 5.0, 5.4, 5.4, 1.4]
        font_size = 6.2
    else:
        widths = [29.0 / columns] * columns
        font_size = 7.5

    header = table.rows[0]
    set_repeat_table_header(header)
    for idx in range(columns):
        cell = header.cells[idx]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, "1F4E79")
        set_cell_border(cell, "B8C6D8")
        set_cell_width(cell, widths[min(idx, len(widths) - 1)])
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(rows[0][idx] if idx < len(rows[0]) else "")
        style_run(run, font_size, bold=True, color="FFFFFF")

    for row_idx, row_values in enumerate(rows[1:], start=1):
        row = table.add_row()
        for idx in range(columns):
            cell = row.cells[idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(cell)
            set_cell_width(cell, widths[min(idx, len(widths) - 1)])
            if row_idx % 2 == 0:
                set_cell_shading(cell, "F7FAFD")
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx in {1, 3, 8} else WD_ALIGN_PARAGRAPH.LEFT
            run = para.add_run(row_values[idx] if idx < len(row_values) else "")
            style_run(run, font_size)

    document.add_paragraph()


def build() -> None:
    text = SOURCE.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(16.54)
    section.page_height = Inches(11.69)
    section.top_margin = Cm(1.15)
    section.bottom_margin = Cm(1.15)
    section.left_margin = Cm(1.0)
    section.right_margin = Cm(1.0)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    styles["Normal"].font.size = Pt(9.5)
    for style_name, size, color in [
        ("Heading 1", 17, "1F4E79"),
        ("Heading 2", 13, "2F5597"),
        ("Heading 3", 10.5, "44546A"),
    ]:
        style = styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)

    add_paragraph(doc, "Báo Cáo Script Unity - Elden Ring", "Title")
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Xuất từ báo cáo phân tích source code ngày 2026-06-09")
    style_run(run, 11, color="44546A")
    doc.add_paragraph()

    i = 1
    while i < len(lines):
        line = lines[i].rstrip()
        if not line:
            i += 1
            continue
        if line.startswith("# "):
            i += 1
            continue
        if line.startswith("## "):
            add_paragraph(doc, line[3:].strip(), "Heading 1")
            i += 1
            continue
        if line.startswith("### "):
            add_paragraph(doc, line[4:].strip(), "Heading 2")
            i += 1
            continue
        if line.startswith("- "):
            para = doc.add_paragraph(style="List Bullet")
            run = para.add_run(strip_inline_markdown(line[2:].strip()))
            style_run(run, 9)
            i += 1
            continue
        if line.startswith("|") and i + 1 < len(lines) and is_separator(lines[i + 1]):
            rows = [split_md_row(line)]
            i += 2
            while i < len(lines) and lines[i].startswith("|"):
                rows.append(split_md_row(lines[i]))
                i += 1
            add_table(doc, rows)
            continue

        para = doc.add_paragraph()
        run = para.add_run(strip_inline_markdown(line))
        style_run(run, 9.5)
        i += 1

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Elden Ring Unity Script Report")
    style_run(run, 8, color="808080")

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
