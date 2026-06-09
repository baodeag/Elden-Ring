from __future__ import annotations

from collections import defaultdict
from dataclasses import dataclass, field
from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SCRIPT_ROOT = ROOT / "Assets" / "Game" / "Scripts"
OUT_MD = ROOT / "outputs" / "game-scripts-detailed-report.md"
OUT_DOCX = ROOT / "outputs" / "game-scripts-detailed-report.docx"


@dataclass
class MethodInfo:
    name: str
    signature: str
    line: int
    body: str = ""
    related: list[str] = field(default_factory=list)
    purpose: str = ""


@dataclass
class ScriptInfo:
    path: Path
    rel: str
    class_name: str
    kind: str
    bases: list[str]
    parents: list[str] = field(default_factory=list)
    children: list[str] = field(default_factory=list)
    refs: list[str] = field(default_factory=list)
    fields: list[str] = field(default_factory=list)
    methods: list[MethodInfo] = field(default_factory=list)
    role: str = ""


def rel(path: Path) -> str:
    return path.relative_to(ROOT).as_posix()


def strip_comments(text: str) -> str:
    text = re.sub(r"/\*[\s\S]*?\*/", " ", text)
    text = re.sub(r"(^|[^:])//.*$", r"\1 ", text, flags=re.MULTILINE)
    return text


def md(value: object) -> str:
    return str(value or "-").replace("|", r"\|").replace("\n", " ")


def short(items: list[str], limit: int = 10) -> str:
    items = list(dict.fromkeys(x for x in items if x))
    if not items:
        return "-"
    suffix = f" +{len(items) - limit}" if len(items) > limit else ""
    return ", ".join(items[:limit]) + suffix


def split_words(name: str) -> str:
    text = re.sub(r"([a-z0-9])([A-Z])", r"\1 \2", name)
    text = text.replace("_", " ")
    return text.strip()


def classify(path: Path, class_name: str, bases: list[str]) -> str:
    low = rel(path).lower()
    if "/editor/" in low:
        return "Editor tool"
    if "NetworkBehaviour" in bases:
        return "NetworkBehaviour"
    if "MonoBehaviour" in bases:
        return "MonoBehaviour"
    if "ScriptableObject" in bases:
        return "ScriptableObject"
    if "StateMachineBehaviour" in bases:
        return "Animator state"
    if any(base.endswith("Effect") for base in bases):
        return "Effect"
    return "Data / plain C#"


def role_for(info: ScriptInfo) -> str:
    low = info.rel.lower()
    name = info.class_name
    if "/editor/" in low:
        return "Công cụ chạy trong Unity Editor để dựng, sửa, kiểm tra hoặc tự động cấu hình scene/prefab/asset."
    if "/world managers/" in low:
        return "Quản lý cấp world: điều phối hệ thống dùng chung như scene, save, object, AI, âm thanh, item hoặc progression."
    if "/character/player/playerui/" in low:
        return "Điều khiển một phần UI của player như HUD, menu, popup, shop, level up, site of grace hoặc equipment."
    if "/character/player/" in low:
        return "Module dành cho người chơi: input, camera, di chuyển, combat, inventory, stats, network hoặc hiệu ứng."
    if "/ai character/states/" in low:
        return "State trong state machine AI, quyết định AI đang idle, truy đuổi, tấn công, ngủ boss hoặc vào combat stance."
    if "/ai character/boss character/" in low:
        return "Module boss AI: quản lý boss, combat, network, phase, skill, collider hoặc âm thanh riêng của boss."
    if "/ai character/knight/" in low:
        return "Kỹ năng/hành vi riêng của Knight boss, gồm projectile, shockwave, slash và VFX."
    if "/ai character/" in low:
        return "Module enemy AI: quản lý hành vi, spawn, animation, locomotion, combat, inventory, sound hoặc network."
    if "/character/" in low:
        return "Base module dùng chung cho Character, được Player và AI kế thừa hoặc tham chiếu."
    if "/colliders/" in low:
        return "Collider/hitbox gây sát thương hoặc phát hiện va chạm để áp dụng damage/effect."
    if "/effects/" in low:
        return "Hiệu ứng tác động lên character: damage tức thời, buff tạm thời, buildup hoặc trạng thái tĩnh."
    if "/function/" in low:
        return "Script chức năng gắn vào object trong scene: interactable, trigger, pickup, elevator, fog wall, dialogue hoặc reset animation flag."
    if "/game saving/" in low:
        return "Dữ liệu/helper phục vụ lưu và tải game."
    if "/items/" in low:
        return "Dữ liệu hoặc hành vi item: weapon, armor, spell, flask, quick slot, projectile hoặc material."
    if "/weapon actions/" in low:
        return "Action ScriptableObject cho vũ khí/spell, được combat manager gọi khi player thực hiện hành động."
    if "/ui/" in low:
        return "UI component độc lập như thanh máu, build-up bar, button effect, slot UI hoặc warning."
    if "/shop/" in low:
        return "Hệ thống shop: tồn kho, entry hàng bán và tương tác mua/bán với player."
    if "/menu scene/" in low:
        return "Điều khiển title/menu scene: load slot, setting menu, preview nhân vật hoặc bắt input menu."
    if "/scenes/" in low:
        return "Load/unload scene hoặc bootstrap các scene world theo location."
    if "/settings/" in low:
        return "Quản lý cấu hình game do người chơi chọn."
    if "/animator/" in low:
        return "Callback trong Animator để reset hoặc toggle parameter animation."
    if "/utility/" in low:
        return "Helper nhỏ dùng chung trong scene hoặc prefab."
    return f"Script gameplay chính cho {split_words(name)}."


def method_purpose(method: MethodInfo, script: ScriptInfo) -> str:
    name = method.name
    readable = split_words(name)
    related = short(method.related, 5)

    exact = {
        "Awake": "Khởi tạo tham chiếu sớm khi object được tạo, thường lấy component con hoặc thiết lập singleton/local cache.",
        "Start": "Chạy sau Awake để hoàn tất thiết lập ban đầu, nhất là khi cần các object khác đã sẵn sàng.",
        "Update": "Cập nhật logic mỗi frame như input, timer, trạng thái chiến đấu hoặc UI.",
        "FixedUpdate": "Cập nhật theo bước vật lý, thường xử lý movement, trigger hoặc physics.",
        "LateUpdate": "Cập nhật cuối frame, thường dùng cho camera, animation hoặc đồng bộ trạng thái sau movement.",
        "OnEnable": "Đăng ký event/callback hoặc bật trạng thái khi component được enable.",
        "OnDisable": "Gỡ event/callback hoặc dọn trạng thái khi component bị disable.",
        "OnDestroy": "Dọn đăng ký/event/tài nguyên khi object bị hủy.",
        "OnValidate": "Tự kiểm tra/cập nhật giá trị trong Unity Editor khi inspector thay đổi.",
        "OnNetworkSpawn": "Thiết lập dữ liệu và callback khi NetworkObject được spawn trong Netcode.",
        "OnNetworkDespawn": "Dọn dữ liệu/callback khi NetworkObject bị despawn khỏi Netcode.",
        "OnTriggerEnter": "Xử lý khi collider khác đi vào trigger của object này.",
        "OnTriggerExit": "Xử lý khi collider khác rời khỏi trigger của object này.",
        "OnCollisionEnter": "Xử lý va chạm vật lý khi object chạm object khác.",
        "OnAnimatorMove": "Nhận root motion từ Animator để áp dụng movement/rotation theo animation.",
    }
    if name in exact:
        base = exact[name]
    elif name.endswith("ServerRpc"):
        base = f"Gửi yêu cầu lên server trong Netcode để server xử lý {split_words(name[:-9]).lower()}."
    elif name.endswith("ClientRpc"):
        base = f"Server gọi xuống client để đồng bộ hoặc phát hiệu ứng cho {split_words(name[:-9]).lower()}."
    elif name.startswith("Handle"):
        base = f"Xử lý luồng {split_words(name[6:]).lower()}."
    elif name.startswith("Try"):
        base = f"Thử thực hiện {split_words(name[3:]).lower()}, thường có kiểm tra điều kiện trước khi chạy."
    elif name.startswith("Attempt"):
        base = f"Cố gắng kích hoạt {split_words(name[7:]).lower()} nếu trạng thái hiện tại cho phép."
    elif name.startswith("Can"):
        base = f"Kiểm tra có được phép {split_words(name[3:]).lower()} hay không."
    elif name.startswith("Is"):
        base = f"Kiểm tra điều kiện/trạng thái {split_words(name[2:]).lower()}."
    elif name.startswith("Set"):
        base = f"Thiết lập giá trị hoặc trạng thái {split_words(name[3:]).lower()}."
    elif name.startswith("Get"):
        base = f"Lấy dữ liệu {split_words(name[3:]).lower()} cho hệ thống khác sử dụng."
    elif name.startswith("Refresh"):
        base = f"Làm mới dữ liệu/hiển thị {split_words(name[7:]).lower()}."
    elif name.startswith("Update"):
        base = f"Cập nhật {split_words(name[6:]).lower()} theo trạng thái mới."
    elif name.startswith("Load"):
        base = f"Nạp dữ liệu hoặc scene liên quan tới {split_words(name[4:]).lower()}."
    elif name.startswith("Save"):
        base = f"Lưu dữ liệu liên quan tới {split_words(name[4:]).lower()}."
    elif name.startswith("Open"):
        base = f"Mở UI/trạng thái/luồng {split_words(name[4:]).lower()}."
    elif name.startswith("Close"):
        base = f"Đóng UI/trạng thái/luồng {split_words(name[5:]).lower()}."
    elif name.startswith("Enable"):
        base = f"Bật {split_words(name[6:]).lower()}."
    elif name.startswith("Disable"):
        base = f"Tắt {split_words(name[7:]).lower()}."
    elif name.startswith("Reset"):
        base = f"Đưa {split_words(name[5:]).lower()} về trạng thái mặc định."
    elif name.startswith("Calculate"):
        base = f"Tính toán {split_words(name[9:]).lower()} từ chỉ số hoặc dữ liệu hiện có."
    elif name.startswith("Apply"):
        base = f"Áp dụng {split_words(name[5:]).lower()} lên character/object mục tiêu."
    elif name.startswith("Take"):
        base = f"Nhận/xử lý tác động {split_words(name[4:]).lower()}, thường là damage hoặc hiệu ứng."
    elif name.startswith("Play"):
        base = f"Phát {split_words(name[4:]).lower()}, thường là animation, sound hoặc VFX."
    elif name.startswith("Spawn") or name.startswith("Create"):
        action = "spawn" if name.startswith("Spawn") else "tạo"
        base = f"{action.capitalize()} object/dữ liệu {split_words(name[5:] if name.startswith('Spawn') else name[6:]).lower()}."
    elif name.startswith("Equip"):
        base = f"Trang bị {split_words(name[5:]).lower()} và cập nhật model/chỉ số liên quan."
    elif name.startswith("Unequip"):
        base = f"Gỡ trang bị {split_words(name[7:]).lower()} và cập nhật lại trạng thái liên quan."
    elif name.startswith("Interact"):
        base = "Thực hiện hành động tương tác khi player chọn object này."
    elif name.startswith("Damage"):
        base = f"Gây hoặc xử lý sát thương cho {split_words(name[6:]).lower()}."
    elif name.startswith("Find"):
        base = f"Tìm {split_words(name[4:]).lower()} trong scene/danh sách dữ liệu."
    elif name.startswith("Add"):
        base = f"Thêm {split_words(name[3:]).lower()} vào danh sách, trạng thái hoặc dữ liệu."
    elif name.startswith("Remove"):
        base = f"Loại bỏ {split_words(name[6:]).lower()} khỏi danh sách, trạng thái hoặc dữ liệu."
    else:
        base = f"Thực hiện logic {readable.lower()} trong script {script.class_name}."

    if method.related:
        return f"{base} Liên kết trực tiếp: {related}."
    return base


def find_matching_brace(text: str, open_index: int) -> int:
    depth = 0
    in_string = False
    quote = ""
    escape = False
    for idx in range(open_index, len(text)):
        ch = text[idx]
        if in_string:
            if escape:
                escape = False
            elif ch == "\\":
                escape = True
            elif ch == quote:
                in_string = False
            continue
        if ch in {'"', "'"}:
            in_string = True
            quote = ch
            continue
        if ch == "{":
            depth += 1
        elif ch == "}":
            depth -= 1
            if depth == 0:
                return idx
    return len(text)


def extract_methods(text: str, class_name: str) -> list[MethodInfo]:
    normal_pattern = re.compile(
        r"(?P<sig>\b(?:public|private|protected|internal)?\s*"
        r"(?:override\s+|virtual\s+|abstract\s+|static\s+|async\s+|new\s+|sealed\s+)*"
        r"[\w<>\[\],\.]+\s+(?P<name>\w+)\s*\([^;{}]*\)\s*(?:\{|=>|;))",
        re.MULTILINE,
    )
    constructor_pattern = re.compile(
        rf"(?P<sig>\b(?:public|private|protected|internal)\s+"
        rf"(?P<name>{re.escape(class_name)})\s*\([^;{{}}]*\)\s*(?:\{{|=>|;))",
        re.MULTILINE,
    )
    methods: list[MethodInfo] = []
    matches = sorted(
        list(normal_pattern.finditer(text)) + list(constructor_pattern.finditer(text)),
        key=lambda item: item.start(),
    )
    seen_starts: set[int] = set()
    for match in matches:
        if match.start() in seen_starts:
            continue
        seen_starts.add(match.start())
        name = match.group("name")
        if name in {"if", "for", "foreach", "while", "switch", "catch", "using", "lock"}:
            continue
        sig = re.sub(r"\s+", " ", match.group("sig")).strip()
        line = text.count("\n", 0, match.start()) + 1
        body = ""
        end = match.end()
        if sig.endswith("{"):
            open_index = text.find("{", match.start(), match.end())
            close_index = find_matching_brace(text, open_index)
            body = text[open_index:close_index + 1]
        elif "=>" in sig:
            semi = text.find(";", match.end())
            body = text[match.start(): semi + 1 if semi != -1 else match.end()]
        methods.append(MethodInfo(name=name, signature=sig.rstrip("{; "), line=line, body=body))
    return methods


def load_infos() -> list[ScriptInfo]:
    files = sorted(SCRIPT_ROOT.rglob("*.cs"))
    declaration_re = re.compile(
        r"\b(?:public|internal|private|protected)?\s*(?:partial\s+)?"
        r"(?:abstract\s+|sealed\s+|static\s+)?(?P<kind>class|struct|interface|enum)\s+"
        r"(?P<name>\w+)(?:\s*:\s*(?P<bases>[^\{\n]+))?"
    )
    field_re = re.compile(
        r"((?:\[SerializeField\]\s*)?(?:public|private|protected|internal)\s+"
        r"(?!class\b|struct\b|enum\b|interface\b|void\b)"
        r"(?:readonly\s+|static\s+|const\s+)?[\w<>\[\],\.]+\s+\w+)\s*(?:=|;)"
    )

    raw: list[tuple[Path, str, str, list[dict[str, object]]]] = []
    all_names: set[str] = set()
    for path in files:
        text = path.read_text(encoding="utf-8", errors="ignore")
        clean = strip_comments(text)
        declarations = []
        for match in declaration_re.finditer(clean):
            bases = [
                part.strip().split()[0]
                for part in (match.group("bases") or "").split(",")
                if part.strip()
            ]
            declaration = {"kind": match.group("kind"), "name": match.group("name"), "bases": bases}
            declarations.append(declaration)
            all_names.add(match.group("name"))
        raw.append((path, text, clean, declarations))

    infos: list[ScriptInfo] = []
    for path, text, clean, declarations in raw:
        primary = next((d for d in declarations if d["kind"] == "class"), declarations[0] if declarations else None)
        class_name = str(primary["name"]) if primary else path.stem
        bases = list(primary["bases"]) if primary else []
        methods = extract_methods(clean, class_name)
        fields = []
        for match in field_re.finditer(clean):
            value = re.sub(r"\s+", " ", match.group(1)).strip()
            if "public " in value or "[SerializeField]" in value:
                fields.append(value)
        refs = sorted(
            name for name in all_names
            if name != class_name and re.search(rf"\b{re.escape(name)}\b", clean)
        )
        info = ScriptInfo(
            path=path,
            rel=rel(path),
            class_name=class_name,
            kind=classify(path, class_name, bases),
            bases=bases,
            refs=refs,
            fields=list(dict.fromkeys(fields)),
            methods=methods,
        )
        info.role = role_for(info)
        infos.append(info)

    by_name = {info.class_name: info for info in infos}
    for info in infos:
        info.parents = [base for base in info.bases if base in by_name]
    for info in infos:
        for parent in info.parents:
            by_name[parent].children.append(info.class_name)
    for info in infos:
        for method in info.methods:
            method.related = sorted(
                name for name in all_names
                if name != info.class_name and method.body and re.search(rf"\b{re.escape(name)}\b", method.body)
            )
            method.purpose = method_purpose(method, info)
    return infos


def build_markdown(infos: list[ScriptInfo]) -> None:
    groups: dict[str, list[ScriptInfo]] = defaultdict(list)
    for info in infos:
        parts = info.rel.split("/")
        group = "/".join(parts[:4]) if len(parts) > 4 else "/".join(parts[:-1])
        if "Character" in parts:
            idx = parts.index("Character")
            group = "/".join(parts[: min(len(parts) - 1, idx + 3)])
        groups[group].append(info)

    lines: list[str] = []
    lines.append("# Báo Cáo Chi Tiết Script Chính - Assets/Game/Scripts")
    lines.append("")
    lines.append("Ngày tạo: 2026-06-09.")
    lines.append(f"Tổng số script phân tích: **{len(infos)}**.")
    lines.append(f"Tổng số hàm/method ghi nhận: **{sum(len(info.methods) for info in infos)}**.")
    lines.append("")
    lines.append("## Cách đọc")
    lines.append("")
    lines.append("- Báo cáo chỉ lấy file `.cs` nằm trong `Assets/Game/Scripts`.")
    lines.append("- `Cha` là class nội bộ dự án mà script kế thừa; `Con` là script kế thừa trực tiếp từ script đó.")
    lines.append("- `Liên kết script` là các class/script nội bộ xuất hiện trong code của script.")
    lines.append("- `Liên kết trong hàm` là các class/script nội bộ xuất hiện trong thân hàm đó.")
    lines.append("- Phần giải thích hàm được viết ngắn, dễ hiểu dựa trên tên hàm, vị trí module, kiểu Unity callback/RPC và các class được hàm dùng.")
    lines.append("")
    lines.append("## Tổng quan module")
    lines.append("")
    lines.append("| Module | Số script | Số hàm | Ý nghĩa |")
    lines.append("|---|---:|---:|---|")
    module_roles = {
        "Assets/Game/Scripts/Character": "Nền nhân vật, Player, AI, boss, combat, network, stats, locomotion và UI player.",
        "Assets/Game/Scripts/Colliders": "Hitbox/collider gây sát thương, parry/block và projectile collision.",
        "Assets/Game/Scripts/Effects": "Hiệu ứng instant/timed/static như damage, buff, buildup, poison, frost, burn.",
        "Assets/Game/Scripts/Function": "Object tương tác/trigger trong scene như pickup, elevator, fog wall, site of grace, dialogue.",
        "Assets/Game/Scripts/Items": "Dữ liệu và hành vi item: weapon, spell, armor, flask, material, quick slot.",
        "Assets/Game/Scripts/Weapon Actions": "Action được weapon/combat gọi để đánh, aim, bắn projectile hoặc cast spell.",
        "Assets/Game/Scripts/World Managers": "Manager cấp world: save, scene, object, AI, sound, item database, boss catalog, progression.",
        "Assets/Game/Scripts/UI": "UI component độc lập như bar, slot, button animation/sound.",
        "Assets/Game/Scripts/Menu Scene": "Title/menu scene, save slots, settings và preview.",
        "Assets/Game/Scripts/Editor": "Công cụ Unity Editor nội bộ để build/fix/setup asset, scene, boss, merchant.",
        "Assets/Game/Scripts/Scenes": "Load scene/additive scene/bootstrap world location.",
        "Assets/Game/Scripts/Shop": "Merchant/shop inventory và mua bán.",
        "Assets/Game/Scripts/Settings": "Game settings.",
        "Assets/Game/Scripts/Animator": "Animator callbacks.",
        "Assets/Game/Scripts/Utility": "Helper nhỏ dùng chung.",
    }
    top_groups = defaultdict(list)
    for info in infos:
        parts = info.rel.split("/")
        key = "/".join(parts[:4])
        top_groups[key].append(info)
    for key in sorted(top_groups):
        items = top_groups[key]
        lines.append(f"| {md(key)} | {len(items)} | {sum(len(i.methods) for i in items)} | {md(module_roles.get(key, 'Module gameplay chính.'))} |")
    lines.append("")
    lines.append("## Quan hệ kế thừa chính")
    lines.append("")
    edges = []
    for info in infos:
        for parent in info.parents:
            edges.append((parent, info.class_name))
    for parent, child in sorted(edges):
        lines.append(f"- `{parent}` -> `{child}`")
    lines.append("")
    lines.append("## Chi tiết từng script")
    lines.append("")
    for group in sorted(groups):
        lines.append(f"### {group}")
        lines.append("")
        for info in sorted(groups[group], key=lambda x: x.rel):
            lines.append(f"#### {info.class_name}")
            lines.append("")
            lines.append(f"- **Đường dẫn:** `{info.rel}`")
            lines.append(f"- **Loại:** {info.kind}")
            lines.append(f"- **Vai trò dễ hiểu:** {info.role}")
            lines.append(f"- **Kế thừa/cha:** {short(info.parents or info.bases, 8)}")
            lines.append(f"- **Script con:** {short(sorted(info.children), 12)}")
            lines.append(f"- **Field public/serialized chính:** {short(info.fields, 12)}")
            lines.append(f"- **Liên kết script:** {short(info.refs, 16)}")
            lines.append("")
            lines.append("| Hàm | Dòng | Ý nghĩa | Liên kết trong hàm |")
            lines.append("|---|---:|---|---|")
            if not info.methods:
                lines.append("| - | - | Script này chủ yếu là dữ liệu/enum/struct, không có method rõ ràng. | - |")
            for method in info.methods:
                lines.append(
                    f"| `{md(method.signature)}` | {method.line} | {md(method.purpose)} | {md(short(method.related, 10))} |"
                )
            lines.append("")
    OUT_MD.write_text("\n".join(lines), encoding="utf-8")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = "D7DEE8") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:color"), color)


def set_cell_width(cell, width_cm: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_cm * 567)))
    tc_w.set(qn("w:type"), "dxa")


def repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def run_style(run, size: float = 8.5, bold: bool = False, color: str | None = None) -> None:
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_text(doc: Document, text: str, style: str | None = None, size: float = 9.5) -> None:
    para = doc.add_paragraph(style=style)
    run = para.add_run(text)
    if style == "Title":
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_style(run, 22, True, "1F4E79")
    elif style in {"Heading 1", "Heading 2", "Heading 3"}:
        run_style(run, size, True)
    else:
        run_style(run, size)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float], font_size: float = 7.0) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    repeat_header(table.rows[0])
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, "1F4E79")
        set_cell_border(cell, "B8C6D8")
        set_cell_width(cell, widths[idx])
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_style(para.add_run(header), font_size, True, "FFFFFF")
    for row_idx, values in enumerate(rows):
        row = table.add_row()
        for idx, value in enumerate(values):
            cell = row.cells[idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(cell)
            set_cell_width(cell, widths[idx])
            if row_idx % 2 == 1:
                set_cell_shading(cell, "F7FAFD")
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx == 1 else WD_ALIGN_PARAGRAPH.LEFT
            run_style(para.add_run(value), font_size)
    doc.add_paragraph()


def build_docx(infos: list[ScriptInfo]) -> None:
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(16.54)
    section.page_height = Inches(11.69)
    section.top_margin = Cm(1.05)
    section.bottom_margin = Cm(1.05)
    section.left_margin = Cm(1.0)
    section.right_margin = Cm(1.0)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    styles["Normal"].font.size = Pt(9)
    for style_name, size, color in [
        ("Heading 1", 16, "1F4E79"),
        ("Heading 2", 12.5, "2F5597"),
        ("Heading 3", 10.5, "44546A"),
    ]:
        style = styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)

    add_text(doc, "Báo Cáo Chi Tiết Script Chính", "Title")
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_style(subtitle.add_run("Phạm vi: Assets/Game/Scripts - giải thích từng script và từng hàm"), 11, False, "44546A")
    doc.add_paragraph()
    add_text(doc, f"Tổng số script: {len(infos)}. Tổng số hàm/method: {sum(len(i.methods) for i in infos)}.", size=10)
    add_text(doc, "Báo cáo này chỉ lấy file .cs trong Assets/Game/Scripts. Các script addon, plugin, TutorialInfo và Assets/Game/System không được đưa vào.", size=9.5)

    add_text(doc, "Tổng Quan Module", "Heading 1")
    top_groups = defaultdict(list)
    module_roles = {
        "Assets/Game/Scripts/Character": "Nền nhân vật, Player, AI, boss, combat, network, stats, locomotion và UI player.",
        "Assets/Game/Scripts/Colliders": "Hitbox/collider gây sát thương, parry/block và projectile collision.",
        "Assets/Game/Scripts/Effects": "Hiệu ứng instant/timed/static như damage, buff, buildup, poison, frost, burn.",
        "Assets/Game/Scripts/Function": "Object tương tác/trigger trong scene như pickup, elevator, fog wall, site of grace, dialogue.",
        "Assets/Game/Scripts/Items": "Dữ liệu và hành vi item: weapon, spell, armor, flask, material, quick slot.",
        "Assets/Game/Scripts/Weapon Actions": "Action được weapon/combat gọi để đánh, aim, bắn projectile hoặc cast spell.",
        "Assets/Game/Scripts/World Managers": "Manager cấp world: save, scene, object, AI, sound, item database, boss catalog, progression.",
        "Assets/Game/Scripts/UI": "UI component độc lập như bar, slot, button animation/sound.",
        "Assets/Game/Scripts/Menu Scene": "Title/menu scene, save slots, settings và preview.",
        "Assets/Game/Scripts/Editor": "Công cụ Unity Editor nội bộ để build/fix/setup asset, scene, boss, merchant.",
        "Assets/Game/Scripts/Scenes": "Load scene/additive scene/bootstrap world location.",
        "Assets/Game/Scripts/Shop": "Merchant/shop inventory và mua bán.",
        "Assets/Game/Scripts/Settings": "Game settings.",
        "Assets/Game/Scripts/Animator": "Animator callbacks.",
        "Assets/Game/Scripts/Utility": "Helper nhỏ dùng chung.",
    }
    for info in infos:
        key = "/".join(info.rel.split("/")[:4])
        top_groups[key].append(info)
    add_table(
        doc,
        ["Module", "Script", "Hàm", "Ý nghĩa"],
        [[key, str(len(items)), str(sum(len(i.methods) for i in items)), module_roles.get(key, "Module gameplay chính.")]
         for key, items in sorted(top_groups.items())],
        [6.0, 2.0, 2.0, 19.0],
        8.0,
    )

    add_text(doc, "Chi Tiết Từng Script", "Heading 1")
    for info in sorted(infos, key=lambda x: x.rel):
        add_text(doc, info.class_name, "Heading 2")
        summary_rows = [
            ["Đường dẫn", info.rel],
            ["Loại", info.kind],
            ["Vai trò", info.role],
            ["Kế thừa/cha", short(info.parents or info.bases, 10)],
            ["Script con", short(sorted(info.children), 14)],
            ["Field chính", short(info.fields, 16)],
            ["Liên kết script", short(info.refs, 20)],
        ]
        add_table(doc, ["Mục", "Nội dung"], summary_rows, [4.5, 24.5], 7.7)
        method_rows = []
        if not info.methods:
            method_rows.append(["-", "-", "Script này chủ yếu là dữ liệu/enum/struct, không có method rõ ràng.", "-"])
        else:
            for method in info.methods:
                method_rows.append([method.signature, str(method.line), method.purpose, short(method.related, 12)])
        add_table(doc, ["Hàm", "Dòng", "Ý nghĩa", "Liên kết trong hàm"], method_rows, [8.0, 1.4, 13.2, 7.4], 6.6)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_style(footer.add_run("Assets/Game/Scripts Detailed Report"), 8, False, "808080")
    doc.save(OUT_DOCX)


def main() -> None:
    infos = load_infos()
    build_markdown(infos)
    build_docx(infos)
    print(f"Markdown: {OUT_MD}")
    print(f"DOCX: {OUT_DOCX}")
    print(f"Scripts: {len(infos)}")
    print(f"Methods: {sum(len(info.methods) for info in infos)}")


if __name__ == "__main__":
    main()
